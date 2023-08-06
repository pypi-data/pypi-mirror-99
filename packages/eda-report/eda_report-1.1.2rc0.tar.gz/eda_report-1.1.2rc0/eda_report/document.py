import logging

from docx import Document
from docx.shared import Inches
from tqdm import tqdm

from eda_report.multivariate import MultiVariable
from eda_report.univariate import Variable


logging.basicConfig(
    format='[%(levelname)s %(asctime)s.%(msecs)03d] %(message)s',
    level=logging.INFO, datefmt='%H:%M:%S'
)


class ReportDocument:
    """The blueprint for objects with methods to create and populate *.docx*
    files.

    The *input data as a whole* is processed as an instance of
    :class:`~eda_report.multivariate.MultiVariable` to get *group statistics
    and plots*. *Individual columns/features* are processed as instances of
    :class:`~eda_report.univariate.Variable` to get statistics and plots for
    *each feature*.

    The `python-docx`_ module is then used to write the results as a *Word*
    document.

    .. _python-docx: https://python-docx.readthedocs.io/en/latest/
    """
    def __init__(self, data, title='Exploratory Data Analysis Report',
                 output_filename='eda-report.docx',
                 graph_color='orangered', table_style='Table Grid'):
        """Initialise an instance of
        :class:`~eda_report.document.ReportDocument`.

        :param data: The data to analyse.
        :type data: Array-like, sequence, iterable
        :param title: The top level heading for the report, defaults to
            'Exploratory Data Analysis Report'.
        :type title: str, optional
        :param output_filename: The name to give the generated report file,
            defaults to 'eda-report.docx'.
        :type output_filename: str, optional
        :param graph_color: A valid matplotlib color specifier, defaults
            to 'orangered'.
        :type graph_color: str, optional
        :param table_style: *Microsoft Word* table style to apply to the
            created tables, defaults to 'Table Grid'.
        :type table_style: str, optional
        """
        self.data = data
        self.TITLE = title
        self.GRAPH_COLOR = graph_color
        self.TABLE_STYLE = table_style
        self.OUTPUT_FILENAME = output_filename
        self.document = Document()
        self._get_report()

    def _get_report(self):
        """Calculate summary statistics, plot graphs, and save the results as
        a .docx file.
        """
        logging.info('Assessing correlation in numeric variables...')
        self.variables = MultiVariable(self.data,
                                       graph_color=self.GRAPH_COLOR)
        logging.info('Done. Summarising each variable...')
        self._create_title_page()  # begin the report document
        self._get_variable_info()  # summarise each variable

        if hasattr(self.variables, 'var_pairs'):
            self._get_bivariate_analysis()  # summarise variable pairs

        self._save_file()
        logging.info(f'Done. Results saved as {self.OUTPUT_FILENAME!r}')

    def _create_title_page(self):
        """Add a title and a brief summary of the data.
        """
        # Title
        self.document.add_heading(self.TITLE, level=0)

        # Brief summary
        num_rows, num_cols = self.data.shape
        num_numeric = self.data.select_dtypes(include='number').columns.size

        if num_rows > 1:
            rows = f'{num_rows} rows (observations)'
        else:
            rows = '1 row (observation)'

        if num_cols > 1:
            cols = f'{num_cols} columns (features)'
        else:
            cols = '1 column (feature)'

        if num_numeric > 1:
            numeric = f', {num_numeric} of which are numeric'
        elif num_numeric == 1:
            numeric = ', 1 of which is numeric'
        else:
            numeric = ''

        intro = f'The dataset consists of {rows} and {cols}{numeric}.'

        self.document.add_paragraph(intro)

        if hasattr(self.variables, 'joint_scatterplot'):
            self.document.add_picture(self.variables.joint_scatterplot,
                                      width=Inches(6.5))
            self.document.add_page_break()

    def _get_variable_info(self):
        """Get a brief introduction, summary statistics, and graphs for each
        individual variable.
        """
        for idx, col in enumerate(tqdm(self.data.columns, ncols=79),
                                  start=1):
            var = Variable(self.data[col], graph_color=self.GRAPH_COLOR)
            # Heading
            self.document.add_heading(f'{idx}. {col}'.title(), level=2)
            # Introduction
            p = self.document.add_paragraph()
            p.add_run(f'{col}'.capitalize()).bold = True
            p.add_run(f' is a {var.var_type} variable ')
            p.add_run(f'with {var.num_unique} unique values. ')
            p.add_run(f'{var.missing} of its values are missing.')
            # Summary statistics
            self.document.add_heading('Summary Statistics', level=4)
            self._create_table(var.statistics)
            if hasattr(var, 'most_common_items'):
                self.document.add_heading('Most Common Values', level=4)
                self._create_table(var.most_common_items)
            # Graphs
            self.document.add_picture(var._graphs, width=Inches(5))

            self.document.add_page_break()

    def _get_bivariate_analysis(self):
        """Get comparisons and scatterplots for pairs of numeric variables.
        """
        self.document.add_heading('Bivariate Analysis (Correlation)', level=1)
        self.document.add_paragraph()
        self.document.add_picture(self.variables.joint_correlation_plot,
                                  width=Inches(6.7))
        self.document.add_page_break()

        for idx, var_pair in enumerate(self.variables.var_pairs, start=1):
            # Heading
            self.document.add_heading(
                f'{idx}. {var_pair[0]} vs {var_pair[1]}'.title(),
                level=2)
            # Introductory text
            p = self.document.add_paragraph()
            p.add_run(f'{var_pair[0]}'.capitalize()).bold = True
            p.add_run(' and ')
            p.add_run(f'{var_pair[1]}'.capitalize()).bold = True
            p.add_run(f' have {self.variables.corr_type[var_pair]}.')
            # Scatter-plot
            self.document.add_picture(
                self.variables.bivariate_scatterplots[var_pair],
                width=Inches(6)
            )
            self.document.add_paragraph()

            if idx % 2 == 0:
                # Add page break after every 2 pairs
                self.document.add_page_break()

    def _create_table(self, data):
        """Create a table from the given data and add it to the document.
        """
        table = self.document.add_table(rows=len(data), cols=2)
        # Set table style
        table.style = self.document.styles[self.TABLE_STYLE]
        # Set column dimensions
        table.columns[0].width = Inches(2.5)
        table.columns[1].width = Inches(2)

        # Get a list of tuples, each to be a row in the table
        items = [(f'{label}', f'{val}') for label, val in data.items()]

        for idx, row in enumerate(table.rows):
            label, value = items[idx]
            row.cells[0].text = label
            row.cells[1].text = value
        self.document.add_paragraph()

    def _save_file(self):
        """Save the document as a .docx file.
        """
        self.document.save(self.OUTPUT_FILENAME)
