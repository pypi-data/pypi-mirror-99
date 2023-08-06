"""
The :py:mod:`vc2_pseudocode_parser.docx_generator` module implements a simplified
wrapper around :py:mod:`docx` for generating Word documents containing SMPTE
specification style code listings.

Document model
--------------

Documents are defined using a hierarchy of the following classes:

.. autoclass:: ListingDocument
    :members:
    :undoc-members:

.. autoclass:: Paragraph
    :members:
    :undoc-members:

.. autoclass:: Run
    :members:
    :undoc-members:

.. autoclass:: RunStyle
    :members:
    :undoc-members:

.. autoclass:: ListingTable
    :members:
    :undoc-members:

.. autoclass:: ListingLine
    :members:
    :undoc-members:

"""

import docx  # type: ignore

import lxml  # type: ignore

from typing import Union, List, Optional

from dataclasses import dataclass, field

from enum import Enum

from copy import deepcopy


def parse_docx_xml(xml: str) -> lxml.etree.Element:
    first_tag_open, first_tag_close, rest = xml.partition(">")
    nsdecls = docx.oxml.ns.nsdecls("w")
    xml = f"{first_tag_open} {nsdecls}{first_tag_close}{rest}"
    return docx.oxml.parse_xml(xml)


PSEUDOCODE_FONT_NAME = "Courier"

LISTING_TABLE_CELL_HMARGIN_MM = 1
LISTING_TABLE_TEXT_VSPACING_MM = 1

LISTING_TABLE_CODE_WIDTH_MM = 166 - 25
LISTING_TABLE_COMMENTS_WIDTH_MM = 25


class TableStyle(Enum):
    """Table styles."""

    pseudocode_listing = "Pseudocode Listing"


class RunStyle(Enum):
    """Text styles for runs."""

    pseudocode = "Pseudocode"
    pseudocode_fdef = "Pseudocode Function Defnition"
    pseudocode_keyword = "Pseudocode Keyword"
    pseudocode_label = "Pseudocode Label"


@dataclass
class ListingDocument:
    """
    A document containing code listings. The root of the document heirarchy.
    """

    body: List[Union["Paragraph", "ListingTable"]] = field(default_factory=list)
    """
    The contents of this document.
    """

    def make_docx_document(self) -> docx.Document:
        """
        Construct a :py:class:`docx.Document` object from this document (ready
        for saving as a file or further manipulation).
        """
        docx_document = docx.Document()

        self._add_run_styles_to_docx_document(docx_document)
        self._add_table_styles_to_docx_document(docx_document)

        for paragraph_or_table in self.body:
            paragraph_or_table._add_to_docx_document(docx_document)

        return docx_document

    def _add_run_styles_to_docx_document(self, docx_document: docx.Document) -> None:
        def make_style(run_style: RunStyle) -> docx.styles.style.BaseStyle:
            style = docx_document.styles.add_style(
                run_style.value,
                docx.enum.style.WD_STYLE_TYPE.CHARACTER,
            )
            style.hidden = False
            style.quick_style = True
            return style

        pseudocode_style = make_style(RunStyle.pseudocode)
        pseudocode_style.font.name = PSEUDOCODE_FONT_NAME

        pseudocode_fdef_style = make_style(RunStyle.pseudocode_fdef)
        pseudocode_fdef_style.base_style = pseudocode_style
        pseudocode_fdef_style.font.bold = True
        pseudocode_fdef_style.font.italic = True

        pseudocode_keyword_style = make_style(RunStyle.pseudocode_keyword)
        pseudocode_keyword_style.base_style = pseudocode_style
        pseudocode_keyword_style.font.bold = True

        pseudocode_label_style = make_style(RunStyle.pseudocode_label)
        pseudocode_label_style.base_style = pseudocode_style
        pseudocode_label_style.font.italic = True

    def _add_table_styles_to_docx_document(self, docx_document: docx.Document) -> None:
        listing_style = docx_document.styles.add_style(
            TableStyle.pseudocode_listing.value,
            docx.enum.style.WD_STYLE_TYPE.TABLE,
        )
        listing_style.hidden = False
        listing_style.quick_style = True
        listing_style.element.append(
            parse_docx_xml(
                """
                <w:tblPr>
                    <!-- Add borders around all cells -->
                    <w:tblBorders>
                        <w:top w:val="single" w:sz="4" w:space="0" w:color="000000" />
                        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000" />
                        <w:start w:val="single" w:sz="4" w:space="0" w:color="000000" />
                        <w:end w:val="single" w:sz="4" w:space="0" w:color="000000" />
                        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000" />
                        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000" />
                    </w:tblBorders>
                    <!-- Set cell margins -->
                    <w:tblCellMar>
                        <w:top w:w="0" w:type="dxa"/>
                        <w:bottom w:w="0" w:type="dxa"/>
                        <w:start w:w="{hmargin}" w:type="dxa"/>
                        <w:end w:w="{hmargin}" w:type="dxa"/>
                    </w:tblCellMar>
                </w:tblPr>
                """.format(
                    hmargin=docx.shared.Mm(LISTING_TABLE_CELL_HMARGIN_MM).twips
                )
            )
        )
        listing_style.element.append(
            parse_docx_xml(
                """
                <w:tcPr>
                    <!-- Vertically align text in cells -->
                    <w:vAlign w:val="center"/>
                </w:tcPr>
                """
            )
        )
        listing_style.element.append(
            parse_docx_xml(
                """
                <w:tblStylePr w:type="firstRow">
                    <!-- Shade first row of table grey -->
                    <w:tcPr>
                        <w:shd w:val="clear" w:color="auto" w:fill="D9D9D9"/>
                    </w:tcPr>
                </w:tblStylePr>
                """
            )
        )
        listing_style.element.append(
            parse_docx_xml(
                """
                <w:tblStylePr w:type="lastCol">
                    <!-- Center text in comment column -->
                    <w:pPr>
                        <w:jc w:val="center"/>
                    </w:pPr>
                </w:tblStylePr>
                """
            )
        )
        listing_style.element.append(
            parse_docx_xml(
                """
                <w:pPr>
                    <!-- Set vertical spacing for text in table -->
                    <w:spacing
                        w:before="{spacing}"
                        w:after="{spacing}"
                        w:beforeAutospacing="0"
                        w:afterAutospacing="0"
                    />
                </w:pPr>
                """.format(
                    spacing=docx.shared.Mm(LISTING_TABLE_TEXT_VSPACING_MM).twips
                )
            )
        )


@dataclass
class Paragraph:
    """
    A paragraph of text consisting of a series of concatenated :py:class:`Runs
    <Run>` of text.

    As a convenience, the constructor accepts either a list of :py:class:`Runs
    <Run>`, a single :py:class:`Run` or a :py:class:`str`.

    :py:class:`Paragraph` objects support the ``+`` operator which will
    concatenate the runs in a pair of :py:class:`Paragraphs <Paragraph>`,
    producing a new :py:class:`Paragraph`. You can also add :py:class:`Run` and
    :py:class:`str` to :py:class:`Paragraphs <Paragraph>` with similar effect.
    """

    runs: List["Run"] = field(default_factory=list)
    """
    The runs of text contained in this paragraph.
    """

    def __init__(self, runs: Union[List["Run"], "Run", str] = []) -> None:
        if isinstance(runs, str):
            self.runs = [Run(runs)]
        elif isinstance(runs, Run):
            self.runs = [runs]
        else:
            self.runs = runs

    def __add__(self, rhs: Union["Paragraph", "Run", str]) -> "Paragraph":
        if isinstance(rhs, Paragraph):
            return Paragraph(self.runs + rhs.runs)
        elif isinstance(rhs, Run):
            return Paragraph(self.runs + [rhs])
        elif isinstance(rhs, str):
            return Paragraph(self.runs + [Run(rhs)])
        else:
            raise TypeError(type(rhs))

    def __radd__(self, lhs: Union["Paragraph", "Run", str]) -> "Paragraph":
        if isinstance(lhs, Paragraph):
            return Paragraph(lhs.runs + self.runs)
        elif isinstance(lhs, Run):
            return Paragraph([lhs] + self.runs)
        elif isinstance(lhs, str):
            return Paragraph([Run(lhs)] + self.runs)
        else:
            raise TypeError(type(lhs))

    def _add_to_docx_document(self, docx_document: docx.Document) -> None:
        docx_paragraph = docx_document.add_paragraph()
        for run in self.runs:
            run.add_to_docx_paragraph(docx_document, docx_paragraph)

    def __str__(self) -> str:
        return "".join(str(run) for run in self.runs)

    def __bool__(self) -> bool:
        return any(self.runs)


@dataclass
class Run:
    """
    A run of text within a paragraph, which will be rendered with a particular
    style (defined in :py:class:`RunStyle`).
    """

    text: str = ""
    style: Optional[RunStyle] = None

    def add_to_docx_paragraph(
        self,
        docx_document: docx.Document,
        docx_paragraph: docx.text.paragraph.Paragraph,
    ) -> None:
        docx_run = docx_paragraph.add_run()
        docx_run.text = self.text
        if self.style is not None:
            docx_run.style = docx_document.styles[self.style.value]

    def __str__(self) -> str:
        return self.text

    def __bool__(self) -> bool:
        return bool(self.text)


@dataclass
class ListingTable:
    """
    A code listing table giving the source code for a single function definiton.

    If any lines in the listing contain a comment, the resulting table will
    have two colums with the code on the left and comments on the right. If no
    lines contain a comment, the table will have a single column.
    """

    rows: List["ListingLine"]
    """The rows in the table."""

    def _add_to_docx_document(self, docx_document: docx.Document) -> None:
        style = docx_document.styles[TableStyle.pseudocode_listing.value]

        has_comments = any(row.comment for row in self.rows)

        docx_table = docx_document.add_table(
            len(self.rows), 2 if has_comments else 1, style.name
        )

        # Force the table column sizes specified below to be used
        docx_table.allow_autofit = False

        # Docx bodge: Table alignment cannot be specified as part of the table style
        docx_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER

        # Docx bodge: Table column widths are not part of a table style so
        # apply these directly here
        code_width = docx.shared.Mm(LISTING_TABLE_CODE_WIDTH_MM)
        comment_width = docx.shared.Mm(LISTING_TABLE_COMMENTS_WIDTH_MM)
        if has_comments:
            docx_table.columns[0].width = code_width
            docx_table.columns[1].width = comment_width
        else:
            docx_table.columns[0].width = code_width + comment_width

        for row, docx_row in zip(self.rows, docx_table.rows):
            code_paragraph = docx_row.cells[0].paragraphs[0]
            for run in row.code.runs:
                run.add_to_docx_paragraph(docx_document, code_paragraph)

            if has_comments:
                comment_paragraph = docx_row.cells[1].paragraphs[0]
                comment_paragraph.paragraph_format.alignment = (
                    docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                )
                for run in row.comment.runs:
                    run.add_to_docx_paragraph(docx_document, comment_paragraph)

            # Word bodge: Word ignores the table column widths and must have
            # column widths defined on every cell
            docx_row.cells[0].width = code_width
            if has_comments:
                docx_row.cells[1].width = comment_width

        # LibreOffice bodge: tblCellMar is ignored in tblPr in the table style
        # and so the tblPr must be copied directly onto the table for it to
        # apply...
        style_tbl_pr = deepcopy(style.element.xpath("w:tblPr")[0])
        docx_table._element.tblPr.extend(style_tbl_pr)


@dataclass
class ListingLine:
    """
    A single row in a :py:class:`ListingTable`.
    """

    code: Paragraph = field(default_factory=Paragraph)
    comment: Paragraph = field(default_factory=Paragraph)
