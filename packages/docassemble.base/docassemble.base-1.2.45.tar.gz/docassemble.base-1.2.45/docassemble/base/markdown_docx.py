import docassemble.base.filter
from docassemble.base.logger import logmessage
import docx.opc.constants
from docxtpl import RichText
from bs4 import BeautifulSoup, NavigableString, Tag
import re

class SoupParser(object):
    def __init__(self, tpl):
        self.paragraphs = [dict(params=dict(style='p', indentation=0), runs=[RichText('')])]
        self.current_paragraph = self.paragraphs[-1]
        self.run = self.current_paragraph['runs'][-1]
        self.bold = False
        self.italic = False
        self.underline = False
        self.strike = False
        self.indentation = 0
        self.style = 'p'
        self.still_new = True
        self.size = None
        self.tpl = tpl
    def new_paragraph(self):
        if self.still_new:
            logmessage("new_paragraph is still new and style is " + self.style + " and indentation is " + str(self.indentation))
            self.current_paragraph['params']['style'] = self.style
            self.current_paragraph['params']['indentation'] = self.indentation
            return
        logmessage("new_paragraph where style is " + self.style + " and indentation is " + str(self.indentation))
        self.current_paragraph = dict(params=dict(style=self.style, indentation=self.indentation), runs=[RichText('')])
        self.paragraphs.append(self.current_paragraph)
        self.run = self.current_paragraph['runs'][-1]
        self.still_new = True
    def __str__(self):
        output = ''
        list_number = 1
        for para in self.paragraphs:
            logmessage("Got a paragraph where style is " + para['params']['style'] + " and indentation is " + str(para['params']['indentation']))
            output += '<w:p><w:pPr><w:pStyle w:val="Normal"/>'
            if para['params']['style'] in ('ul', 'ol', 'blockquote'):
                output += '<w:ind w:left="' + str(36*para['params']['indentation']) + '" w:right="0" w:hanging="0"/>'
            output += '<w:rPr></w:rPr></w:pPr>'
            if para['params']['style'] == 'ul':
                output += str(RichText("•\t"))
            if para['params']['style'] == 'ol':
                output += str(RichText(str(list_number) + ".\t"))
                list_number += 1
            else:
                list_number = 1
            for run in para['runs']:
                output += str(run)
            output += '</w:p>'
        return output
    def start_link(self, url):
        ref = self.tpl.docx._part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        self.current_paragraph['runs'].append('<w:hyperlink r:id="%s">' % (ref, ))
        self.new_run()        
        self.still_new = False
    def end_link(self):
        self.current_paragraph['runs'].append('</w:hyperlink>')
        self.new_run()
        self.still_new = False
    def new_run(self):
        self.current_paragraph['runs'].append(RichText(''))
        self.run = self.current_paragraph['runs'][-1]
    def line_break(self):
        self.current_paragraph['runs'].append(RichText("\n"))
        self.run = self.current_paragraph['runs'][-1]
    def traverse(self, elem):
        for part in elem.contents:
            if isinstance(part, NavigableString):
                self.run.add(str(part), italic=self.italic, bold=self.bold, underline=self.underline, strike=self.strike, size=self.size)
                self.still_new = False
            elif isinstance(part, Tag):
                logmessage("Part name is " + str(part.name))
                if part.name == 'p':
                    self.new_paragraph()
                    self.traverse(part)
                elif part.name == 'li':
                    self.new_paragraph()
                    self.traverse(part)
                elif part.name == 'ul':
                    logmessage("Entering a UL")
                    oldstyle = self.style
                    self.style = 'ul'
                    self.indentation += 10
                    self.traverse(part)
                    self.indentation -= 10
                    self.style = oldstyle
                    logmessage("Leaving a UL")
                elif part.name == 'ol':
                    logmessage("Entering a OL")
                    oldstyle = self.style
                    self.style = 'ol'
                    self.indentation += 10
                    self.traverse(part)
                    self.indentation -= 10
                    self.style = oldstyle
                    logmessage("Leaving a OL")
                elif part.name == 'br':
                    self.line_break()
                elif part.name == 'strong':
                    self.bold = True
                    self.traverse(part)
                    self.bold = False
                elif part.name == 'em':
                    self.italic = True
                    self.traverse(part)
                    self.italic = False
                elif part.name == 'strike':
                    self.strike = True
                    self.traverse(part)
                    self.strike = False
                elif part.name == 'u':
                    self.underline = True
                    self.traverse(part)
                    self.underline = False
                elif part.name == 'blockquote':
                    oldstyle = self.style
                    self.style = 'blockquote'
                    self.indentation += 20
                    self.traverse(part)
                    self.indentation -= 20
                    self.style = oldstyle
                elif re.match(r'h[1-6]', part.name):
                    oldsize = self.size
                    self.size = 60 - ((int(part.name[1]) - 1) * 10)
                    self.new_paragraph()
                    self.bold = True
                    self.traverse(part)
                    self.bold = False
                    self.size = oldsize
                elif part.name == 'a':
                    self.start_link(part['href'])
                    self.underline = True
                    self.traverse(part)
                    self.underline = False
                    self.end_link()
            else:
                logmessage("Encountered a " + part.__class__.__name__)

def markdown_to_docx(text, tpl):
    logmessage('Starting conversion of ' + text)
    source = docassemble.base.filter.markdown_to_html(text, do_terms=False)
    source = re.sub("\n", ' ', source)
    source = re.sub(">\s+<", '><', source)
    soup = BeautifulSoup('<html>' + source + '</html>', 'html.parser')
    logmessage('Now it has become ' + source)
    parser = SoupParser(tpl)
    for elem in soup.find_all(recursive=False):
        parser.traverse(elem)
    output = str(parser)
    logmessage("Pre-final output is: " + output)
    output = re.sub(r'\[EMOJI ([^,\]]+), *([0-9A-Za-z.%]+)\]', lambda x: docassemble.base.filter.image_include_docx_template(x), output)
    output = re.sub(r'\[FILE ([^,\]]+), *([0-9A-Za-z.%]+), *([^\]]*)\]', lambda x: docassemble.base.filter.image_include_docx_template(x), output)
    output = re.sub(r'\[FILE ([^,\]]+), *([0-9A-Za-z.%]+)\]', lambda x: docassemble.base.filter.image_include_docx_template(x), output)
    output = re.sub(r'\[FILE ([^,\]]+)\]', lambda x: docassemble.base.filter.image_include_docx_template(x), output)
    output = re.sub(r'\[QR ([^,\]]+), *([0-9A-Za-z.%]+), *([^\]]*)\]', docassemble.base.filter.qr_include_docx_template, output)
    output = re.sub(r'\[QR ([^,\]]+), *([0-9A-Za-z.%]+)\]', docassemble.base.filter.qr_include_docx_template, output)
    output = re.sub(r'\[QR ([^\]]+)\]', docassemble.base.filter.qr_include_docx_template, output)
    logmessage("Final output is: " + output)
    return output
