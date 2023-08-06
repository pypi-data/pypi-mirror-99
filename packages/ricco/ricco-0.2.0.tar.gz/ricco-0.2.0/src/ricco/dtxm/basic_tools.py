# encoding:GBK
import os

import matplotlib.pyplot as plt
import pandas as pd
from ricco import rdf
from ricco.dtxm.basic_docx import _Docx
from ricco.dtxm.wiki import class_dic
from ricco.util import col_round
from ricco.util import ensure_list

from docx.shared import Inches
from docx.shared import RGBColor


class _Tools(_Docx):

    def __init__(self, filename, cols_list: list = None):
        super().__init__()
        self.filename = filename

        if isinstance(self.filename, str):
            self.add_intense_quote(f'Document��{self.filename}')
            self.df = rdf(self.filename)
        elif isinstance(self.filename, pd.DataFrame):
            self.df = self.filename
        else:
            raise ValueError('������Dataframe��·��')
        if cols_list != None:
            cols_list = ensure_list(cols_list)
            self.df = self.df[[cols_list]]
        self.length = len(self.df)

    def serise_describe(self, col):
        '''��ֵ���е�������ͳ��'''
        desc = pd.DataFrame(self.df[col].describe().reset_index())
        desc = desc.rename(columns={'index': '����', col: 'ֵ'})
        desc['����'] = desc['����'].replace(to_replace=class_dic)
        skew_add = pd.DataFrame({'����': 'ƫ��ϵ��', 'ֵ': [self.df[col].skew()]})
        kurt_add = pd.DataFrame({'����': '���ϵ��', 'ֵ': [self.df[col].kurt() - 3]})
        null_num = self.length - desc.loc[desc['����'] == '����', 'ֵ'][0]
        null_rate = null_num / self.length
        null_add = pd.DataFrame({'����': 'ȱʧ��', 'ֵ': [null_num]})
        null_rate_add = pd.DataFrame({'����': 'ȱʧ��', 'ֵ': [null_rate]})
        desc = desc.append(skew_add, sort=False)
        desc = desc.append(kurt_add, sort=False)
        desc = desc.append(null_add, sort=False)
        desc = desc.append(null_rate_add, sort=False)
        desc = col_round(desc, 'ֵ')
        self.add_df2table(desc)

    def object_describe(self, col):
        desc = pd.DataFrame(self.df[col].value_counts().reset_index())
        if len(desc) > 20:
            desc = desc.rename(columns={'index': '����_Top15', col: '����'})
            desc['����_Top15'] = desc['����_Top15'].replace(to_replace=class_dic)
            res = desc.head(15)
        else:
            desc = desc.rename(columns={'index': '����', col: '����'})
            desc['����'] = desc['����'].replace(to_replace=class_dic)
            res = desc
        self.add_df2table(res)

    def is_float(self, col):
        def try2float(x):
            try:
                return float(x)
            except ValueError:
                return None

        length = len(self.df[~self.df[col].isna()])
        null_num = len(self.df[~self.df[col].apply(lambda x: try2float(x)).isna()])
        rates = null_num / length
        if rates >= 0.8:
            text = f' {col} ���г���{int(rates * 100)}%��ֵΪ��ֵ�͵�����'
            p = self.doc.add_paragraph('')
            p.add_run(text).font.color.rgb = RGBColor(250, 0, 0)

    def is_date(self, col):
        length = len(self.df[~self.df[col].isna()])
        null_num = len(self.df[~pd.to_datetime(self.df[col], errors='coerce').isna()])
        rates = null_num / length
        if rates >= 0.8:
            text = f' {col} ���г���{int(rates * 100)}%��ֵΪ���ڸ�ʽ������'
            p = self.doc.add_paragraph('')
            p.add_run(text).font.color.rgb = RGBColor(250, 0, 0)

    def hist_plot(self, col):
        plt.figure(figsize=(12, 4))
        plt.style.use('seaborn')
        data = self.df[~self.df[col].isna()][col].values
        plt.hist(data)
        plt.savefig('image.png')
        self.doc.add_picture('image.png', width=Inches(6))
        if os.path.exists('image.png'):
            os.remove('image.png')
