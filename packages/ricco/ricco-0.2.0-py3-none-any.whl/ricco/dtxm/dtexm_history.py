# encoding:GBK
from docx import Document
from docx.shared import Inches
from ricco import rdf
import pandas as pd
import numpy as np
import os
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from ricco.util import fn
from ricco.dtxm.wiki import class_dic
from ricco.util import col_round
from ricco.util import ensure_list
import matplotlib.pyplot as plt


class Dtexm(object):
    def __init__(self, filename, cols_list: list = None):
        self.filename = filename
        self.doc = Document()
        # ����Ĭ������
        self.doc.styles['Normal'].font.name = u'΢���ź�'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'΢���ź�')
        self.doc.styles['Normal'].font.size = Pt(10.5)
        self.doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
        self.doc.add_heading('���ݼ�ⱨ��', 0)

        if isinstance(self.filename, str):
            self.add_intense_quote(f'Document��{self.filename}')
            self.df = rdf(self.filename)
        elif isinstance(self.filename, pd.DataFrame):
            self.df = self.filename
        else:
            raise ValueError('������Dataframe��·��')
        if cols_list != None:
            cols_list = ensure_list(cols_list)
            self.df = self.df[[cols_list]]
        self.lenth = len(self.df)

    def add_df2table(self, table_df):
        '''��dataframe��Ϊ���д��'''
        t = self.doc.add_table(table_df.shape[0] + 1, table_df.shape[1], style='Table Grid')
        for j in range(table_df.shape[-1]):
            t.cell(0, j).text = table_df.columns[j]
        for i in range(table_df.shape[0]):
            for j in range(table_df.shape[-1]):
                t.cell(i + 1, j).text = str(table_df.values[i, j])

    def add_normal_p(self, text):
        '''�����ͨ����'''
        self.doc.add_paragraph(text)

    def add_bullet_list(self, text):
        '''��������б�'''
        self.doc.add_paragraph(text, style='List Bullet')

    def add_order_list(self, text):
        '''��������б�'''
        self.doc.add_paragraph(text, style='List Number')

    def add_intense_quote(self, text):
        self.doc.add_paragraph(text, style='Intense Quote')

    def add_title(self, text, n):
        '''��ӱ���'''
        self.doc.add_heading(text, level=n)

    def serise_describe(self, col):
        '''��ֵ���е�������ͳ��'''
        desc = pd.DataFrame(self.df[col].describe().reset_index())
        desc = desc.rename(columns={'index': '����', col: 'ֵ'})
        desc['����'] = desc['����'].replace(to_replace=class_dic)
        skew_add = pd.DataFrame({'����': 'ƫ��ϵ��',
                                 'ֵ': [self.df[col].skew()]})
        kurt_add = pd.DataFrame({'����': '���ϵ��',
                                 'ֵ': [self.df[col].kurt() - 3]})

        null_num = self.lenth - desc.loc[desc['����'] == '����', 'ֵ'][0]
        null_rate = null_num / self.lenth

        null_add = pd.DataFrame({'����': 'ȱʧ��',
                                 'ֵ': [null_num]})
        null_rate_add = pd.DataFrame({'����': 'ȱʧ��',
                                      'ֵ': [null_rate]})
        desc = desc.append(skew_add, sort=False)
        desc = desc.append(kurt_add, sort=False)
        desc = desc.append(null_add, sort=False)
        desc = desc.append(null_rate_add, sort=False)
        desc = col_round(desc, 'ֵ')
        self.add_df2table(desc)

    def object_describe(self, col):
        desc = pd.DataFrame(self.df[col].value_counts().reset_index())
        if len(desc) > 20:
            desc = desc.rename(columns={'index': '����_Top15', col: '����'})
            desc['����_Top15'] = desc['����_Top15'].replace(to_replace=class_dic)
            res = desc.head(15)
        else:
            desc = desc.rename(columns={'index': '����', col: '����'})
            desc['����'] = desc['����'].replace(to_replace=class_dic)
            res = desc
        self.add_df2table(res)

    def is_float(self, col):
        def try2float(x):
            try:
                return float(x)
            except ValueError:
                return None

        length = len(self.df[~self.df[col].isna()])
        null_num = len(self.df[~self.df[col].apply(lambda x: try2float(x)).isna()])
        rates = null_num / length
        if rates >= 0.8:
            text = f' {col} ���г���{int(rates * 100)}%��ֵΪ��ֵ�͵�����'
            p = self.doc.add_paragraph('')
            p.add_run(text).font.color.rgb = RGBColor(250, 0, 0)

    def is_date(self, col):
        length = len(self.df[~self.df[col].isna()])
        null_num = len(self.df[~pd.to_datetime(self.df[col], errors='coerce').isna()])
        rates = null_num / length
        if rates >= 0.8:
            text = f' {col} ���г���{int(rates * 100)}%��ֵΪ���ڸ�ʽ������'
            p = self.doc.add_paragraph('')
            p.add_run(text).font.color.rgb = RGBColor(250, 0, 0)

    def hist_plot(self, col):
        plt.figure(figsize=(12, 4))
        plt.style.use('seaborn')
        data = self.df[~self.df[col].isna()][col].values
        plt.hist(data)
        plt.savefig('image.png')
        self.doc.add_picture('image.png', width=Inches(6))
        if os.path.exists('image.png'):
            os.remove('image.png')

    def basic(self):
        '''���ݻ�����Ϣ����'''
        # ����
        self.add_bullet_list('������')
        self.add_normal_p('��'.join(self.df.columns))
        # �ļ�size
        self.add_bullet_list('�ļ��ߴ磺')
        self.add_normal_p(f'�У�{self.df.shape[0]}����{self.df.shape[1]}')
        # ������
        self.add_bullet_list('�������ͣ�')
        self.col_types = pd.DataFrame(self.df.dtypes, columns=['����']).reset_index().rename(columns={'index': '����'})
        self.add_df2table(self.col_types)

    def col_by_col(self):
        '''���м��'''
        for col in self.df.columns:
            self.add_bullet_list(col)
            if (self.df[col].dtype == 'int64') | (self.df[col].dtype == 'float64'):
                self.serise_describe(col)
                self.hist_plot(col)
                self.add_normal_p('')
            else:
                self.is_float(col)
                self.is_date(col)
                self.add_normal_p('')

    def save(self, savefilename: str = None):
        '''�����ļ���word�ĵ�'''
        if savefilename == None:
            if isinstance(self.filename, pd.DataFrame):
                raise FileNotFoundError('������Ҫ������ļ�·��')
            savefilename = fn(self.filename) + '-' '��ⱨ��.docx'
        self.doc.save(savefilename)
        print('�ļ�������', os.path.abspath(savefilename))

    def examine_all(self):
        '''��������'''
        self.basic()
        self.col_by_col()
        self.save()


if __name__ == '__main__':
    doc = Dtexm('sample.csv')

    doc.examine_all()
