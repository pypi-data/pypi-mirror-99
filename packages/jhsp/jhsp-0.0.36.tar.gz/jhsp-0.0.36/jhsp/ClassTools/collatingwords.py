import docx
import os
import sys
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

class CollatingWords():
    """
    This class, named 'collating words' , is built to collate some words from JH's doc through
    removing duplication, merging similar unit and adjusting the sequence of units.

    The basic operating unit of this class is unit, composed by some words which are similar in
    meaning or spelling.
    """

    def __init__(self,old_dir):
        self.old_dir = old_dir
        self.old_doc = docx.Document(old_dir)
        self.units = None
        self.get_units()  #从docx里读取unit

    def get_units(self):
        """
        从当前类的docx对象中读取units的函数，没有返回值，直接赋值到类属性。
        :return:
        """

        # units ,with a nature of list,is used to store unit
        units = []

        # units ,with a nature of dict,is used to store words and its meanings.
        unit = {}

        # unit是按段落分开的，一个段落就是一个unit
        for paragraph in self.old_doc.paragraphs:

            # unit与unit之间有一个空段落，因此只要读到‘’ or ‘ ’ 且之前读到过具体内容，就是一个unit结束了
            if paragraph.text in ['', ' ']:
                if unit:   #检验是否之前读到过具体内容，如果是，则将这个unit添加进units列表里面
                    unit['index'] = list(unit.keys())      #为了方便后续操作，将unit的key也就是单词做成一个index。
                    units.append(unit.copy())
                unit = {}

            else:   #如果读到的不是空段落，则证明读到的是一个unit 里的其中一行。  一行就是一个单词+ ‘ ’ +解释

                if 'Unit:' in paragraph.text : continue  # 跳过标题行

                word_meaning = paragraph.text.strip().split(' ', 1)  # 用空格分隔开单词和解释
                word = word_meaning[0].lower()       # 将单词统一转化为小写

                if word == '':continue              # 如果单词是‘ ’ 则证明这是一个无效的单词，跳过当前循环

                try:                                # 如果存在单词，但不存在解释，则自动补充解释为 待补充
                    meaning = word_meaning[1].strip()
                except IndexError:
                    meaning = '待补充'

                unit[word] = meaning               # 讲这个单词极其解释以key-value的形式加入到当前unit中


        self.units = units

    def save_cdoc(self):
        """
        这个函数是用来保存整理之后的units到doc的，没有返回值，保存路径与读取路径一致。
        """

        new_doc = docx.Document()

        for i, unit in enumerate(self.units):

            new_doc.add_paragraph('Unit:' + str(i))
            for key, value in unit.items():
                if key == 'index': continue
                new_doc.add_paragraph(key + ' '*5 + value)
            new_doc.add_paragraph()


        new_dir = os.path.dirname(self.old_dir)
        new_dir = os.path.join(new_dir,'new.docx')

        new_doc.save(new_dir)

    def rp(self):
        """
        this function, completely named removing duplication, is built to removing some lists
        which are identical to one of forward lists
        """

        dup_unit_index = []       # 要删除的重复unit的索引列表
        unit_words = []           # 将unit的index的值（一个包含该unit所有单词的列表）储存到一个列表

        for index, unit in enumerate(self.units):     # 循环遍历所有unit
            if unit['index'] not in unit_words:       # 判断当前unit的index是否已经存在，如果存在就是重复的unit
                unit_words.append(unit['index'])      # 如果没有不存在，则没有重复，将这个index加入到列表中。
            else:
                dup_unit_index.append(index)         # 将重复unit的索引添加进要删除的重复unit的索引列表

        dup_unit_index.sort(reverse=True)            # 逆序排列，这样从后往前按索引删除列表元素，删除一个之后，其余的要删除元素的索引不会改变。
        for i in dup_unit_index:
            print('删除重复unit：{}'.format(self.units[i]['index']))       # 打印要删除的元素
            del self.units[i]                        # 删除元素

    def rs(self):

        """
        这个函数是 removing single word。如果一个unit只有一个单词，那没什么意义，直接删除。
        """
        single_unit_index = []

        for index, unit in enumerate(self.units):
            if len(unit['index']) <=1 :                   # 如果index 的长度<=1,证明该unit最多只含有一个单词
                single_unit_index.append(index)

        single_unit_index.sort(reverse=True)              # 删除方法与去重一样。
        for i in single_unit_index:
            print('删除仅有一个单词的unit：{}'.format(self.units[i]['index']))
            del self.units[i]

    def sort_by_corr(self):
        """
        将相关的unit组成一个unit集合，按不同集合的unit数量进行逆序排序，然后赋值到类units中。（方便记忆的时候区块化记忆）
        :return:
        """

        def get_similar_uid(ori_uid):
            """
            这个函数在units 中找出与参数unit相关的unit
            参数：目标unit的索引id
            返回值：相关unit的索引id列表，每一个元素就是一个相关unit的id

            """

            similar_uid = []

            for uid, unit in enumerate(self.units):      # 遍历unit

                if uid == ori_uid: continue              # 如果当前遍历的unit就是目标unit，则跳过

                # 这个求相似的逻辑，现在是按照单词完全一样，后期可以改成一个函数，只要有60%子字符串一致就算相似。
                if len(set(self.units[uid]['index']).intersection(set(self.units[ori_uid]['index']))) >= 1:
                    similar_uid.append(uid)     # 如果相似，则将当前遍历的unit的索引id加入到相似列表

            return similar_uid

        similar_units_list = []        # 相似unit组成的units列表为元素，组成的一个列表。列表元素是相似units列表，相似units列表的元素是unit
        similar_uid_list = []          # 当前彼此相似unit的id索引列表

        for cur_id, cur_unit in enumerate(self.units):  # 遍历units

            next = 0                                    # 与后续的next代码一起实现，如果当前unit已经存在于某个相似units之中，则跳过当前循环
            for similar_units in similar_units_list:    # 从相似units列表中遍历相似units
                for unit in similar_units:              # 从相似units中遍历unit
                    if cur_unit == unit: next = 1       # 如果当前unit 已经存在于某个相似units之中，则跳过当前循环

            if next == 1: continue


            similar_uid_list.append(cur_id)                    # 如果当前unit不存在于任何一个相似units，则添加到一个新相似units中
            similar_uid_list.extend(get_similar_uid(cur_id))   # 将与这unit相关的unit添加到这个相似units中

            i = len(similar_uid_list) - 1                      # 设置一个阈值i，如果经过一次相关寻找相关unit之后，这个相似units中
                                                               # unit的数量并没有增加，则证明所有与之相关的unit都找出来了，则结束循环。

            while len(similar_uid_list) > i :                  # 如果相似units 大于阈值，也就是改变了，则进入循环
                i = len(similar_uid_list)                      # 首先赋值阈值等于units的长度，这样如果units的长度没有改变，则会结束循环

                temp_similar_uid = similar_uid_list.copy()     # 复制一个当前units，这样可以在遍历循环的同时，进行改变
                for uid in temp_similar_uid :                  # 遍历当前相似units 中每一个unit
                    similar_uid_list.extend(get_similar_uid(uid)) # 找出每一个unit的相似unit并加入到这个当前units中
                    similar_uid_list = list(set(similar_uid_list)) # 这样会有重复，因此取集合以排除重复。

            temp_similar_units = [self.units[i] for i in similar_uid_list]  # 通过uid，把这个相似units中的所有unit添加到一个暂时units列表中
            similar_units_list.append(temp_similar_units)    # 把这个units的暂时列表当成一个元素添加到相似units列表中
            similar_uid_list = []

        similar_units_list = sorted(similar_units_list,key=lambda x: len(x), reverse=True)     # 逆序排序

        temp_similar_units = []      # 赋值给类的units
        for units in similar_units_list:
            temp_similar_units.extend(units)
        self.units = temp_similar_units

    def run(self):
        self.rp()
        self.rs()
        self.sort_by_corr()
        self.save_cdoc()

