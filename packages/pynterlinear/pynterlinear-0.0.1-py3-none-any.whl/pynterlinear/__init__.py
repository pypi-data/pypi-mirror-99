import re
from docx import Document, shared
import pathlib
import os
import logging
# logging.basicConfig(filename='pynterlinear.log',level=logging.DEBUG)
# logging.disable(logging.CRITICAL)
#These symbols could potentially be used to split up morphemes.
#Some of them are standard, some not.
delimiters = ["-", "–", ".","=",":","*","~","<",">","[","]","(",")","/","\\"]

#Create a hash of common glossing abbreviations and their meaning.
glossing_abbrevs = {}
fn = pathlib.Path(__file__).parent / "./glossing.txt"
raw_glosses = open(fn,"r").read()
for entry in raw_glosses.split("\n"):
    glossing_abbrevs[entry.split("\t")[0]] = entry.split("\t")[1]

#This is used for keeping track of abbreviations that came up in the analysis
#but which are not in glossing_abbrevs.
#This script is primarily geared towards the LaTeX package glossingtool.sty,
#which has the same predefined glossing abbreviations.
unknown_abbrevs = []

def get_unknown_abbrevs():
    return unknown_abbrevs

def get_delims():
    return delimiters

def get_all_abbrevs():
    return glossing_abbrevs

#Can be used to quickly look up the proper abbreviation for something.
def search_abbrev(term):
    output = []
    for key, abbrev in glossing_abbrevs.items():
        if term in abbrev:
            output.append("%s\t%s" % (key, abbrev))
    return "\n".join(output)

#This takes an object word and a gloss and returns a form--meaning hash
def get_morphemes_from_word(object_word,gloss_word):
    morpheme_delimiters = ["-", "=", "~"]
    for delim in morpheme_delimiters:
        object_word = object_word.replace(delim," ")
    for delim in morpheme_delimiters:
        gloss_word = gloss_word.replace(delim," ")
    forms = object_word.split(" ")
    glosses = gloss_word.split(" ")
    if len(forms) != len(glosses):
        print("%s: mismatch in morpheme number!" % object_word)
    return dict(zip(forms, glosses))

#This takes a whole example (hash) and returns a form--meaning hash
def get_morphemes(**example):
    pairings = {}
    if len(example["obj"]) != len(example["gloss"]):
        print("%s: mismatch in word number! %s vs %s words" % (
            example["id"],
            len(example["obj"]),
            len(example["gloss"]))
            )
    for i, word in enumerate(example["obj"]):
        new_entries = get_morphemes_from_word(
            example["obj"][i],example["gloss"][i]
            )
        pairings = {**pairings, **new_entries}
    return pairings
    
#Splits a word up into morphemes and delimiters        
def split_word(word):
    output = []
    char_list = list(word)
    for i, char in enumerate(char_list):
        if len(output) == 0 or (char in delimiters or output[-1] in delimiters):
            output.append(char)
        else:
            output[-1]+=char
    return output

#Takes an uppercase string like 1SG and breaks it up into known abbreviations like ["1", "SG"]
def get_glossing_combination(input):
    output = []
    temp_text = ""
    for i, char in enumerate(list(input)):
        if re.match(r"[1-3]+", char):
            if (i < len(input)-1 and input[i+1]=="+"):
                temp_text+=char
            elif input[i-1]=="+":
                temp_text+=char
                output.append(temp_text)
                temp_text=""
            else:
                if temp_text != "": output.append(temp_text)
                output.append(char)
                temp_text=""
        else:
            temp_text+=char
    if temp_text != "": output.append(temp_text)
    for morpheme in output:
        if (morpheme.lower() not in glossing_abbrevs 
                and morpheme.lower() not in unknown_abbrevs):
            unknown_abbrevs.append(morpheme.lower())
    return output        

#Creates expex code with \gl{} for glossing abbreviations
def get_expex_code(input):
    input = input.replace("\\","\\textbackslash()")
    words_list = input.split(" ")
    #iterate through words of glossing line
    for i, word in enumerate(words_list):
        #Transform _X_ to subscript beforehand
        word = re.sub(r"_([^AP_]+?)_", "\\\\textsubscript{\g<1>}", word)
        output = " "
        #take proper nouns into account
        if (len(word)==2
                and word[0] == word[0].upper()
                and word[1] == "."):
            output+=word
        else:
            parts=split_word(word)
            for j, part in enumerate(parts):
                if (part == part.upper()
                        and part not in delimiters
                        and part != "?"
                        and not (
                            len(part) == 1 
                            and not re.match(r"\d", part) #is it only one capital letter?
                            and (len(parts) == j+1 #are we at the end of the word?
                                or ( #or are there more characters?
                                    parts[j+1] in ["."] #and is the next character a period?
                                    and (j == len(parts)-2 or parts[j+2] in ["-", "="]) #and is that period at the end of the word or followed by a delimiter?
                                )
                            )
                        )
                    ):
                    if part.lower() in glossing_abbrevs:
                        output+= "\\gl{%s}" % part.lower()
                    #take care of numbered genders
                    elif part[0] == "G" and re.match(r"\d", part[1:]):
                        output += "\\gl{g}%s" % (part[1:])
                    else:
                        for extracted_morpheme in get_glossing_combination(
                            part
                        ):
                            output += "\\gl{%s}" % extracted_morpheme.lower()
                else:
                    output+=part
        words_list[i] = output[1:]
    gloss_text_upcased=" ".join(words_list)
    gloss_text_upcased=gloss_text_upcased.replace("~","\\glosstilde{}").replace("_","\_").replace("\_a\_", "_a_").replace("\_p\_", "_p_").replace("textbackslash()", "textbackslash{}")
    return gloss_text_upcased

#This function takes a list of dictionaries and produces ExPex code
#Use for_beamer if you want to insert a \glottolink command
#For this, a key glottocode is needed in the example hash.
def convert_to_expex(examples, for_beamer = False, from_corpus=False, pextag="NEWTAG", latex_labels=True, multicols=False, no_surf=False):
    #See what languages and sources we're dealing with and whether it makes
    #sense to just print them once
    languages = []
    sources = []
    lengths = []
    for example in examples:
        if "language" in example.keys():
            languages.append(example["language"])
        if "source" in example.keys():
            sources.append(example["source"].split("[")[0])
        lengths.append(len(example["obj"])+5)
        lengths.append(len(example["gloss"])+5)
        lengths.append(len(example["trans"]))
    #If there is only one language, we only need to print it once
    if len(list(set(languages))) < 2:
        same_language = True
    else:
        same_language = False
    #If there is only one source, we only need to print it once
    if len(list(set(sources))) == 1:
        same_source = True
    else:
        same_source = False
    #If there is only one example, we just use \ex, and don't insert additional
    #\a.  Otherwise, we use \pex and \a
    #output is where the final text is being built
    if len(examples) == 1:
        pex = False
        output = "\\ex<%s>" % examples[0]["id"]
    else:
        pex = True
        output = "\\pex<%s>" % pextag
    len_threshold = max(lengths)
    if pex and multicols and len_threshold < 35 and same_language and same_source:
        multicols = True
    else:
        multicols = False
    #If there is one and only language specified, we print it right at the
    #start, after \(p)ex
    if same_language and "language" in examples[0].keys():
        #Print \glottolink if for_beamer is set to True
        if for_beamer and 0 == 1:
            language_string = "\\glottolink{%s}{%s}" % (
                examples[0]["glottocode"],
                examples[0]["language"]
            )
        elif same_language and from_corpus:
            language_string = ""
        else:
            language_string = examples[0]["language"]
        if "parnote" in examples[0].keys():
            parnote = examples[0]["parnote"]
        else:
            parnote = ""
        #Add the language name to the output
        output += language_string
        #Same logic as before, if there is only one source (for the one
        #language), we add it only once, at the start
        if same_source:
            logging.debug("All subexamples of pex %s have the same source!" % pextag)
            logging.debug(examples)
            #Gather the page (ranges) given in the single sources in one list
            page_string = []
            for example in examples:
                page_string.append(example["source"].split("[")[1].split("]")[0])
            if len(list(set(page_string))) == 1:
                page_string = page_string[0]
            else:
                page_string = ", ".join(page_string)
            #Print the list joined by commas
            if sources[0] != "pc":
                output += " \\parencite[%s][%s]{%s}" % (parnote, page_string, sources[0])
            else:
                output += " \\perscomm{%s}" % page_string
        if latex_labels and pex:
            output += r"\exl{%s}" % (pextag)
        #beamer for some reason wants a line break after \pex, but absolutely
        #not after \ex; text documents want the opposite
        if (for_beamer and pex) or (not for_beamer and not pex and language_string != ""):
            output += "\\\\"
    #If we're dealing with subexamples, we need a linebreak before the \a
    if pex:     
        output += "\n"
    #Do we want multicols?
        if multicols:
            output += "\\begin{multicols}{2}\n"
    #part_text holds everything between \a (if present) and \endgl.  It is used
    #even in the case of a single example
    for example in examples:
        #We simply add nothing if we're dealing with a single example,
        #otherwise we add an \a with a tag
        if pex:
            part_text = "\\a<%s>" % example["id"]
        else:
            part_text = ""
        #If there is a source, we extract the details
        if "source" in example.keys():
            pages = example["source"].split("[")[1].split("]")[0]
            source_key = example["source"].split("[")[0]
            if "parnote" in example.keys():
                parnote = example["parnote"]
            else:
                parnote = ""
        if from_corpus:
            speaker = ""
            start = ""
            end = ""
            part = ""
            text_id = ""
            if "speaker" in example.keys():
                speaker = example["speaker"]
            if "start" in example.keys():
                start = example["start"]
            if "end" in example.keys():
                end = example["end"]
            if "part" in example.keys():
                part = example["part"]
            if "text_id" in example.keys():
                text_id = example["text_id"]
        #If we're dealing with different languages, we want to print every one,
        #after the respective \a
        if not same_language:
            #Again with the glottolink string
            if for_beamer and 0==1:
                language_string = "\\glottolink{%s}{%s}" % (example["glottocode"], example["language"])
            else:
                language_string = example["language"]
            part_text += " " + language_string
            #And if we're not dealing with a single source, we add all the
            #sources seperately
            if not same_source and not from_corpus:
                if source_key != "pc":
                    part_text += " \\parencite[%s][%s]{%s}" % (parnote, pages, source_key)
                else:
                    part_text += " \\perscomm{%s}" % pages
        #Text documents also want a linebreak after \a with text behind it,
        #beamer doesn't
        if not for_beamer and not same_language:
            part_text += "\\\\"
        #Line break before \begingl
        part_text += "\n"
        #If there is a surface form, we add it, otherwise it's just an empty
        #string
        if "surface" not in example or example["surface"] == example["obj"] or no_surf:# or "-" in example["surface"]:
            surface_string = ""
        else:
            if "[" in example["surface"] and "]" in example["surface"]:
                surface_string = "\\glpreamble \\normalfont %s//\n" % (example["surface"])
            else:
                surface_string = "\\glpreamble %s//\n" % (example["surface"])
        #Get the \gl{}-ified interlinear gloss
        gloss_text_markup = get_expex_code(example["gloss"])
        #Put together the interlinear text
        part_text += "\\begingl\n%s\\gla %s//\n\\glb %s//\n\\glft \\qu{%s}" % (
            surface_string,
            example["obj"].replace("~","\\glosstilde{}"),
            gloss_text_markup,
            example["trans"]
        )
        #If we're dealing with one language (which we printed after \pex), but
        #with different sources, we wanna print the sources after the
        #translation
        if same_language and not same_source and "source" in example.keys():
            if source_key != "pc":
                part_text += " \\parencite[%s][%s]{%s}" % (parnote, pages, source_key)
            else:
                part_text += " \\perscomm{%s}" % pages
        elif from_corpus:
            part_text += " \\textref[speaker=%s, start=%s, end=%s, part=%s]{%s}" % (speaker, start, end, part, text_id)
        # else:
 #            part_text += " \\parencite[%s]{%s}" % (pages, source_key)
        #Finish interlinear text
        part_text += "//"
        if latex_labels and pex:
            part_text += r"\exl{%s.%s}" % (pextag, example["id"])
        elif latex_labels and not pex:
            part_text += r"\exl{%s}" % (example["id"])
        part_text += "\n\\endgl"
        #Add to output
        output += part_text + "\n"
    if multicols:
        output += "\\end{multicols}\n"
    #Finish \ex…\xe
    output += "\\xe"
    return(output)

def convert_to_word(examples, use_tables=True, filename="csv2word_export.docx"):
    
    def get_running_number_tables(document):
        for i in range(1,len(document.tables)+1):
            topright = document.tables[-i].rows[0].cells[0].text
            search = re.search('\((.*)\)', topright)
            if search:
                return int(search.group(1))
        return 0
        
    def get_running_number_tabs(document):
        for i in range(1,len(document.paragraphs)+1):
            partext = document.paragraphs[-i].text
            search = re.search('\((.*)\)', partext)
            if search:
                return int(search.group(1))
        return 0
        
    if filename in os.listdir("."):
        document = Document(filename)
        if use_tables: 
            running_number = get_running_number_tables(document)+1
        else:
            running_number = get_running_number_tabs(document)+1
    else:
        document = Document()
        running_number = 1
    
    if use_tables:
        xs = 1
        if len(examples) > 1:
            xs+=1
        max_length = 0
        exhe = 3
        for exno, example in enumerate(examples):
            obj_words = example["obj"].split(" ")
            gloss_words = example["gloss"].split(" ")
            trans = example["trans"]
            table = document.add_table(rows=exhe, cols=len(obj_words)+xs)
            #Only for development purposes
            # table.style = "Table Grid"
            for i, obj_word in enumerate(obj_words):
                table.rows[0].cells[i+xs].paragraphs[0].add_run(obj_word).italic = True
            for i, gloss_word in enumerate(gloss_words):
                if (len(gloss_word)==2
                        and gloss_word[0] == gloss_word[0].upper()
                        and gloss_word[1] == "."):
                    tt = table.rows[1].cells[i+xs].paragraphs[0].add_run(
                        "%s" % gloss_word
                    )
                else:
                    morphemes = split_word(gloss_word)
                    for morpheme in morphemes:
                        #take proper nouns into account
                        if (morpheme == morpheme.upper()
                                and morpheme not in delimiters
                                and morpheme != "?"):
                            tt = table.rows[1].cells[i+xs].paragraphs[0].add_run(
                                "%s" % morpheme.lower()
                            )
                            tt.font.small_caps = True
                        else:
                            tt = table.rows[1].cells[i+xs].paragraphs[0].add_run(
                                "%s" % morpheme
                            )
            for column in table.columns:
                for cell in column.cells:
                    cell._tc.tcPr.tcW.type = "auto"
            transcell = table.rows[-1].cells[0+xs].merge(table.rows[-1].cells[-1])
            transcell.text = "‘%s’" % trans
            table.rows[0].cells[0].allow_autofit = False 
            table.rows[0].cells[0].width = shared.Cm(1)
            if exno == 0:
                table.rows[0].cells[0].text = f"({running_number})"
            if len(examples) > 1:
                table.rows[0].cells[1].text = "%s." % chr(97+exno)
            #Remove random 10pt spacing in EVERY CELL
            for row in table.rows:
                for cell in row.cells:
                    for par in cell.paragraphs:
                        par.paragraph_format.space_after = shared.Pt(0)
            par = document.add_paragraph("")
            par.paragraph_format.space_after = shared.Pt(0)
    else:
        gloss_para = document.add_paragraph("(%s)" % running_number)
        for exno, example in enumerate(examples):
            if len(examples) > 1:
                sub_ex_label = chr(97+exno)  + ".\t"
                tabstops = [1, 2]
            else:
                sub_ex_label = ""
                tabstops = [1.25]
            gloss_para.add_run("\t%s" % sub_ex_label)
            obj_run = gloss_para.add_run("%s\n\t\t" % (example["obj"].replace(" ", "\t")))
            obj_run.italic = True
            for i, gloss_word in enumerate(example["gloss"].split(" ")):
                if (len(gloss_word)==2
                        and gloss_word[0] == gloss_word[0].upper()
                        and gloss_word[1] == "."):
                    add_text = gloss_para.add_run(gloss_word)
                else:
                    morphemes = split_word(gloss_word)
                    for morpheme in morphemes:
                        #take proper nouns into account
                        if (morpheme == morpheme.upper()
                                and morpheme not in delimiters
                                and morpheme != "?"):
                            add_text = gloss_para.add_run(morpheme.lower())
                            add_text.font.small_caps = True
                        else:
                            add_text = gloss_para.add_run(morpheme)
                if i < len(example["gloss"].split(" "))-1:
                    gloss_para.add_run("\t")    
            gloss_para.add_run("\n\t\t‘%s’" % (example["trans"]))#, sub_example["source"]))
            tab_stops = gloss_para.paragraph_format.tab_stops
            for tabstop in tabstops:
                tab_stops.add_tab_stop(shared.Cm(tabstop))
            gloss_para = document.add_paragraph()
    document.save(filename)